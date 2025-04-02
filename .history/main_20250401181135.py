from docx import Document


def replace_text_in_doc(doc_path, file_path, old_text, new_text):
    doc = Document(doc_path)

    # Paragraflar ichidagi matnni almashtirish
    for para in doc.paragraphs:
        if old_text in para.text:
            para.text = para.text.replace(old_text, new_text)

    # Jadval ichidagi matnni almashtirish
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if old_text in cell.text:
                    cell.text = cell.text.replace(old_text, new_text)

    # Sarlavha va pastki qismni tekshirish
    sections = doc.sections
    for section in sections:
        # Header (sarlavha)
        for para in section.header.paragraphs:
            if old_text in para.text:
                para.text = para.text.replace(old_text, new_text)

        # Footer (pastki qism)
        for para in section.footer.paragraphs:
            if old_text in para.text:
                para.text = para.text.replace(old_text, new_text)

    # O‘zgarishlarni saqlash
    doc.save(file_path)
    print(f"'{old_text}' so‘zi '{new_text}' ga almashtirildi va '{file_path}' fayl saqlandi!")


# Foydalanish:
replace_text_in_doc("example.docx", "357.docx", "2025_raqam", "100")
replace_text_in_doc("357.docx", "357.docx", "2024_raqam", "75")
replace_text_in_doc("357.docx", "357.docx", "2023_raqam", "50")
