from docx import Document
import openpyxl


def from_excel(yacheyka):

    # Load the Excel file
    workbook = openpyxl.load_workbook('baza.xlsx', data_only=True)

    # Select the active sheet or a specific sheet
    # or workbook['SheetName'] if you know the sheet name
    sheet = workbook.active

    # Get the value of cell B4
    value_b4 = sheet[yacheyka].value
    workbook.close()
    value_b4_comma = str(value_b4).replace('.', ',')

    # Print the value
    print(f"The value in cell B4 is: {value_b4_comma}")
    return str(value_b4).replace('.', ',') if value_b4 else ''


def replace_text_in_doc(doc_path, file_path, old_text, new_text):
    try:
        doc = Document(doc_path)

        for para in doc.paragraphs:
            for run in para.runs:
                run.text = run.text.replace(old_text, new_text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.text = run.text.replace(old_text, new_text)

        for section in doc.sections:
            for para in section.header.paragraphs + section.footer.paragraphs:
                for run in para.runs:
                    run.text = run.text.replace(old_text, new_text)

        doc.save(file_path)
        print(
            f"'{old_text}' -> '{new_text}' almashtirildi va '{file_path}' saqlandi!")
    except Exception as e:
        print(f"Xatolik yuz berdi: {e}")


# period
replace_text_in_doc("example.docx", "357.docx", "period_uz", "mart")
replace_text_in_doc("357.docx", "357.docx", "period_ru", "март")


# Foydalanish:
replace_text_in_doc("357.docx", "357.docx", "2025_raqam", "100")
replace_text_in_doc("357.docx", "357.docx", "2023_raqam", from_excel('B4'))
replace_text_in_doc("357.docx", "357.docx", "2024_raqam", from_excel('C4'))


def eng_kup_hudud(column, start_row, end_row, new_keys):
    workbook = openpyxl.load_workbook(file_path, data_only=True)
    sheet = workbook.active

    data_dict = {}

    # Excel qatorlarini yangi kalitlar bilan moslashtirish
    for i, row in enumerate(range(start_row, end_row + 1)):
        if i >= len(new_keys):  # Agar kalitlar ro‘yxati tugagan bo‘lsa, boshqa kalitni ishlatmasin
            break
        cell_value = sheet[f"{column}{row}"].value
        if cell_value is not None:
            try:
                numeric_value = float(str(cell_value).replace(
                    ',', '.'))  # Son shakliga o'tkazish
                data_dict[new_keys[i]] = numeric_value
            except ValueError:
                # Agar noto‘g‘ri qiymat bo‘lsa, 0 qo‘yamiz
                data_dict[new_keys[i]] = 0

    workbook.close()

    # Qiymatlar bo‘yicha Z-A (kamayish tartibida) saralash
    sorted_dict = dict(
        sorted(data_dict.items(), key=lambda item: item[1], reverse=True))

    return sorted_dict


# Yangi kalitlar ro‘yxati (agar kerak bo‘lsa, uzunligini o‘zgartirish)
new_keys = ['Andijon viloyati', 'Buxoro viloyati', 'Jizzax viloyati', 'Qashqadaryo viloyati', 'Navoiy viloyati', 'Namangan viloyati',
            'Samarqand viloyati', 'Surxondaryo viloyati', 'Sirdaryo viloyati', 'Toshkent shahri', 'Toshkent viloyati', "Farg'ona viloyati", 'Xorazm viloyati', "Qoraqalpog'iston Respublikasi"]

# Foydalanish
sorted_data = eng_kup_hudud("B", 5, 18, new_keys)
# print(sorted_data)


def eng_katta(data, val, urin):
    urin = urin - 1
    first_key = list(data.keys())[urin]
    first_value = data[first_key]
    if val == 1:
        return first_key
    else:
        return first_value
