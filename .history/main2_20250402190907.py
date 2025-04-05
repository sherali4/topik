import openpyxl
from docx import Document
import re
from transliterate import translit


def from_excel(yacheyka, workbook):
    sheet = workbook.active
    value = sheet[yacheyka].value
    return str(value).replace('.', ',') if value else ''


def replace_text_in_doc(doc_path, old_text, new_text):
    try:
        doc = Document(doc_path)

        for para in doc.paragraphs:
            para.text = para.text.replace(old_text, new_text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell.text = cell.text.replace(old_text, new_text)

        doc.save(doc_path)
        print(f"'{old_text}' -> '{new_text}'")
    except Exception as e:
        print(f"Xatolik yuz berdi: {e}")


def eng_kup_hudud(column, start_row, end_row, new_keys, sheet):
    data_dict = {new_keys[i]: str(sheet[f"{column}{row}"].value).replace('.', ',')
                 for i, row in enumerate(range(start_row, end_row + 1)) if i < len(new_keys)}
    return dict(sorted(data_dict.items(), key=lambda item: item[1], reverse=True))


def eng_katta(data, val, urin):
    key = list(data.keys())[urin - 1]
    return key if val == 1 else data[key]


def kiril_to_latin(name):
    return translit(name, 'ru', reversed=True).upper() if re.search(r'[а-яА-Я]', name) else name.upper()


class Engkup:
    hududlar = ['Andijon viloyati', 'Buxoro viloyati', 'Jizzax viloyati', 'Qashqadaryo viloyati', 'Navoiy viloyati', 'Namangan viloyati',
                'Samarqand viloyati', 'Surxondaryo viloyati', 'Sirdaryo viloyati', 'Toshkent shahri', 'Toshkent viloyati', "Farg'ona viloyati", 'Xorazm viloyati', "Qoraqalpog'iston Respublikasi"]

    hududlar_ru = {
        'Andijon viloyati': 'Андижанская область',
        'Buxoro viloyati': 'Бухарская область',
        'Jizzax viloyati': 'Джизакская область',
        'Qashqadaryo viloyati': 'Кашкадарьинская область',
        'Navoiy viloyati': 'Навоийская область',
        'Namangan viloyati': 'Наманганская область',
        'Samarqand viloyati': 'Самаркандская область',
        'Surxondaryo viloyati': 'Сурхандарьинская область',
        'Sirdaryo viloyati': 'Сирдарьинская область',
        'Toshkent shahri': 'Город Ташкент',
        'Toshkent viloyati': 'Ташкент область',
        "Farg'ona viloyati": "Ферганская область",
        'Xorazm viloyati': 'Хорезмская область',
        "Qoraqalpog'iston Respublikasi": "Республика Каракалпакстан"
    }

    def __init__(self, file_nomi, uzgarmas, ustun, urin, sheet):
        self.file_nomi = file_nomi
        self.uzgarmas = uzgarmas
        self.ustun = kiril_to_latin(ustun)
        self.urin = urin
        self.sheet = sheet

    def update_doc(self, ru=False):
        data = eng_kup_hudud(self.ustun, 5, 18, self.hududlar, self.sheet)
        kursatkich = eng_katta(data, 1, self.urin)
        kursatkich = self.hududlar_ru[kursatkich] if ru else kursatkich
        replace_text_in_doc(self.file_nomi, self.uzgarmas, kursatkich)


# Excel faylni ochish
workbook = openpyxl.load_workbook('baza.xlsx', data_only=True)
sheet = workbook.active

# Hujjatni yangilash
replace_text_in_doc("example.docx", "period_uz", "mart")
replace_text_in_doc("example.docx", "period_ru", "март")
replace_text_in_doc("example.docx", "2023_raqam", from_excel('B4', sheet))
replace_text_in_doc("example.docx", "2024_raqam", from_excel('C4', sheet))

# Engkup obyektlari bilan ishlash
Engkup('example.docx', '@h1', 'C', 1, sheet).update_doc()
Engkup('example.docx', '@hr1', 'C', 1, sheet).update_doc(ru=True)
Engkup('example.docx', '@h2', 'C', 2, sheet).update_doc()
Engkup('example.docx', '@hr2', 'C', 2, sheet).update_doc(ru=True)

workbook.close()
