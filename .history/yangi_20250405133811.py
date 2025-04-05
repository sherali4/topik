from functions.classes import Filltopic
from docx import Document
# qurilish
qurilish = Filltopic('example.docx', '400.docx', 'sanoat', 'ulush')
qurilish.matn('period_uz', 'mart')
qurilish.matn('period_ru', 'март')
qurilish.respublika('2023_raqam', '2023')
qurilish.respublika('2024_raqam', '2024')
qurilish.respublika('2025_raqam', '2025')
qurilish.kursatkich('@k1', 1)
qurilish.kursatkich('@k2', 2)
qurilish.hudud('@h1', 1)
qurilish.hudud('@h2', 2)
qurilish.hudud_ru('@hr1', 1)
qurilish.hudud_ru('@hr2', 2)


# run.add_break(WD_BREAK.LINE)


# Load your template document
template_doc = Document("shablon/qurilish_1.docx")

from docx import Document

# Load your template document
doc = Document("template.docx")

target_text = "@hudud"

for para in doc.paragraphs:
    for run in para.runs:
        if target_text in run.text:
            print(f"Found text: {run.text}")
            print("  Bold:", run.bold)
            print("  Italic:", run.italic)
            print("  Underline:", run.underline)
            print("  Font name:", run.font.name)
            print("  Font size:", run.font.size)
            print("  Style name:", para.style.name)

