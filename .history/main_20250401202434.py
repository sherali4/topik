from docx import Document
import openpyxl


def from_excel(yacheyka):

    # Load the Excel file
    workbook = openpyxl.load_workbook('baza.xlsx', data_only=True)

    # Select the active sheet or a specific sheet
    # or workbook['SheetName'] if you know the sheet name
    sheet = workbook.active

    # Get the value of cell B4
    value_b4 = sheet[yacheyka].value
    workbook.close()
    value_b4_comma = str(value_b4).replace('.', ',')

    # Print the value
    print(f"The value in cell B4 is: {value_b4_comma}")
    return str(value_b4).replace('.', ',') if value_b4 else ''


def replace_text_in_doc(doc_path, file_path, old_text, new_text):
    try:
        doc = Document(doc_path)

        for para in doc.paragraphs:
            for run in para.runs:
                run.text = run.text.replace(old_text, new_text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.text = run.text.replace(old_text, new_text)

        for section in doc.sections:
            for para in section.header.paragraphs + section.footer.paragraphs:
                for run in para.runs:
                    run.text = run.text.replace(old_text, new_text)

        doc.save(file_path)
        print(
            f"'{old_text}' -> '{new_text}' almashtirildi va '{file_path}' saqlandi!")
    except Exception as e:
        print(f"Xatolik yuz berdi: {e}")


# period
replace_text_in_doc("example.docx", "357.docx", "period_uz", "mart")
replace_text_in_doc("357.docx", "357.docx", "period_ru", "март")


# Foydalanish:
replace_text_in_doc("357.docx", "357.docx", "2025_raqam", "100")
replace_text_in_doc("357.docx", "357.docx", "2023_raqam", from_excel('B4'))
replace_text_in_doc("357.docx", "357.docx", "2024_raqam", from_excel('C4'))


def get_sorted_excel_data_with_custom_keys(file_path, column, start_row, end_row, new_keys):
    workbook = openpyxl.load_workbook(file_path, data_only=True)
    sheet = workbook.active

    data_list = []

    for row in range(start_row, end_row + 1):
        cell_value = sheet[f"{column}{row}"].value
        if cell_value is not None:
            try:
                numeric_value = float(str(cell_value).replace(
                    ',', '.'))  # Son shakliga o'tkazish
                data_list.append(numeric_value)
            except ValueError:
                data_list.append(0)  # Agar noto‘g‘ri qiymat bo‘lsa, 0 qo‘yamiz

    workbook.close()

    # Qiymatlar bo‘yicha saralash (o‘sish tartibida)
    sorted_values = sorted(data_list)

    # Yangi kalitlarni berilgan tartibda moslashtirish
    sorted_dict = {new_keys[i]: sorted_values[i] for i in range(len(new_keys))}

    return sorted_dict


# Yangi kalitlar ro‘yxati
new_keys = [1700, 1703, 1706, 1708, 1710, 1712, 1714,
            1718, 1722, 1724, 1726, 1727, 1730, 1733, 1735]

# Foydalanish
sorted_data_with_new_keys = get_sorted_excel_data_with_custom_keys(
    "baza.xlsx", "B", 4, 18, new_keys)
print(sorted_data_with_new_keys)
