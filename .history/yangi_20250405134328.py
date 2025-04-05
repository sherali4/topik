from functions.classes import Filltopic
from docx import Document
# qurilish
qurilish = Filltopic('example.docx', '400.docx', 'sanoat', 'ulush')
qurilish.matn('period_uz', 'mart')
qurilish.matn('period_ru', 'март')
qurilish.respublika('2023_raqam', '2023')
qurilish.respublika('2024_raqam', '2024')
qurilish.respublika('2025_raqam', '2025')
qurilish.kursatkich('@k1', 1)
qurilish.kursatkich('@k2', 2)
qurilish.hudud('@h1', 1)
qurilish.hudud('@h2', 2)
qurilish.hudud_ru('@hr1', 1)
qurilish.hudud_ru('@hr2', 2)


# run.add_break(WD_BREAK.LINE)


# Load your template document
qurilish = Filltopic('qurilish.docx', '400.docx', 'sanoat', 'ulush')
doc = Document("shablon/qurilish_1.docx")


source_style = None  # To store the style from "@hudud"

# Step 1: Find "@hudud" and save its style
for para in doc.paragraphs:
    for run in para.runs:
        if "@hudud" in run.text:
            source_style = {
                "bold": run.bold,
                "italic": run.italic,
                "underline": run.underline,
                "font_name": run.font.name,
                "font_size": run.font.size,
                "font_color": run.font.color.rgb if run.font.color else None
            }
            break
    if source_style:
        break

# Step 2: Find "kursatkich" and apply the saved style
if source_style:
    for para in doc.paragraphs:
        for run in para.runs:
            if "kursatkich" in run.text:
                run.bold = source_style["bold"]
                run.italic = source_style["italic"]
                run.underline = source_style["underline"]
                run.font.name = source_style["font_name"]
                run.font.size = source_style["font_size"]
                if source_style["font_color"]:
                    run.font.color.rgb = source_style["font_color"]

    doc.save("styled_output.docx")
    print("Style applied successfully!")
else:
    print("Text '@hudud' not found.")
