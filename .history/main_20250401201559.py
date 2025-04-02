from docx import Document
import openpyxl


def from_excel(yacheyka):

    # Load the Excel file
    workbook = openpyxl.load_workbook('baza.xlsx', data_only=True)

    # Select the active sheet or a specific sheet
    # or workbook['SheetName'] if you know the sheet name
    sheet = workbook.active

    # Get the value of cell B4
    value_b4 = sheet[yacheyka].value
    workbook.close()
    value_b4_comma = str(value_b4).replace('.', ',')

    # Print the value
    print(f"The value in cell B4 is: {value_b4_comma}")
    return str(value_b4).replace('.', ',') if value_b4 else ''


def replace_text_in_doc(doc_path, file_path, old_text, new_text):
    try:
        doc = Document(doc_path)

        for para in doc.paragraphs:
            for run in para.runs:
                run.text = run.text.replace(old_text, new_text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.text = run.text.replace(old_text, new_text)

        for section in doc.sections:
            for para in section.header.paragraphs + section.footer.paragraphs:
                for run in para.runs:
                    run.text = run.text.replace(old_text, new_text)

        doc.save(file_path)
        print(
            f"'{old_text}' -> '{new_text}' almashtirildi va '{file_path}' saqlandi!")
    except Exception as e:
        print(f"Xatolik yuz berdi: {e}")


# period
replace_text_in_doc("example.docx", "357.docx", "period_uz", "mart")
replace_text_in_doc("357.docx", "357.docx", "period_ru", "март")


# Foydalanish:
replace_text_in_doc("357.docx", "357.docx", "2025_raqam", "100")
replace_text_in_doc("357.docx", "357.docx", "2023_raqam", from_excel('B4'))
replace_text_in_doc("357.docx", "357.docx", "2024_raqam", from_excel('C4'))


def get_sorted_excel_data_as_dict_desc(file_path, column, start_row, end_row):
    workbook = openpyxl.load_workbook(
        file_path, data_only=True)  # Formulalarni hisoblash
    sheet = workbook.active

    data_dict = {}

    for row in range(start_row, end_row + 1):
        cell_value = sheet[f"{column}{row}"].value
        data_dict[row] = str(cell_value).replace(
            '.', ',') if cell_value else ''

    workbook.close()  # Faylni yopish

    # Teskari tartibda saralangan lug‘at qaytarish
    return dict(sorted(data_dict.items(), reverse=True))


# Foydalanish
sorted_excel_data_desc = get_sorted_excel_data_as_dict_desc(
    "baza.xlsx", "B", 4, 18)
print(sorted_excel_data_desc)



