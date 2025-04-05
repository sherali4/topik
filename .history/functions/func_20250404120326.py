from lugat import tarmoqlar
from openpyxl.utils import get_column_letter
from func import replace_text_in_doc, eng_katta, eng_kup_hudud, kiril_to_latin
import openpyxl
from docx import Document
import re
from transliterate import translit
from lugat import viloyatlar


def from_excel(kitob, yacheyka):
    # Load the Excel file
    workbook = openpyxl.load_workbook('baza.xlsx', data_only=True)
    # sheet = workbook.active
    sheet = workbook[kitob]

    # Get the value of the specified cell
    value = sheet[yacheyka].value
    workbook.close()
    value = str(value).replace('.', ',')
    return value


def replace_text_in_doc(doc_path, file_path, old_text, new_text):
    try:
        doc = Document(doc_path)

        for para in doc.paragraphs:
            for run in para.runs:
                run.text = run.text.replace(old_text, new_text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.text = run.text.replace(old_text, new_text)

        for section in doc.sections:
            for para in section.header.paragraphs + section.footer.paragraphs:
                for run in para.runs:
                    run.text = run.text.replace(old_text, new_text)

        doc.save(file_path)
        print(
            f"'{old_text}' -> '{new_text}'")
    except Exception as e:
        print(f"Xatolik yuz berdi: {e}")


def eng_kup_hudud(column, start_row, end_row, new_keys, kitob):
    workbook = openpyxl.load_workbook('baza.xlsx', data_only=True)
    sheet = workbook[kitob]

    data_dict = {}

    # Excel qatorlarini yangi kalitlar bilan moslashtirish
    for i, row in enumerate(range(start_row, end_row + 1)):
        if i >= len(new_keys):  # Agar kalitlar ro‘yxati tugagan bo‘lsa, boshqa kalitni ishlatmasin
            break
        cell_value = sheet[f"{column}{row}"].value

        try:
            numeric_value = str(cell_value).replace('.', ',')
            data_dict[new_keys[i]] = numeric_value
        except ValueError:
            data_dict[new_keys[i]] = 0

    workbook.close()

    # Qiymatlar bo‘yicha Z-A (kamayish tartibida) saralash
    sorted_dict = dict(
        sorted(data_dict.items(), key=lambda item: item[1], reverse=True))

    return sorted_dict


new_keys = viloyatlar.viloyatlar_uz


def eng_katta(data, val, urin):
    urin = urin - 1
    first_key = list(data.keys())[urin]
    first_value = data[first_key]
    if val == 1:
        return first_key
    else:
        return first_value


def eng_kup(ustun, dan, gacha, new_keys, urin, kitob, hudud=0):
    eng_katta(eng_kup_hudud(ustun, dan, gacha, new_keys, kitob), hudud, urin)


def kiril_to_latin(name):
    if isinstance(name[0], str) and re.search(r'[а-яА-Я]', name[0]):
        latin_name = translit(name, 'ru', reversed=True).upper()
        return latin_name
    else:
        return name.upper()


class Engkup:

    def __init__(self, uzgarmas, urin):
        self.uzgarmas = uzgarmas
        self.urin = urin
        self.kitob = 'ulush'
        self.dan = 5
        self.gacha = self.dan + 13
        self.hududlar = viloyatlar.viloyatlar_uz
        self.hududlar_ru = viloyatlar.viloyatlar_ru

    @property
    def ustun(self):
        return self.ustun

    @property
    def file_nomi(self, value):
        self.file_nomi = value
        return self.file_nomi

    @ustun.setter
    def ustun(self, value):
        self.ustun = kiril_to_latin(value)  # Transliterate to Latin

    @property
    def hudud(self):
        kursatkich = eng_kup_hudud(
            self.ustun, self.dan, self.gacha, self.hududlar, self.kitob)
        kursatkich2 = eng_katta(kursatkich, 1, urin=self.urin)
        replace_text_in_doc(self.file_nomi, self.file_nomi,
                            self.uzgarmas, kursatkich2)

    @property
    def kitob(self, kitob):
        self.kitob = kitob

    @property
    def hudud_ru(self):
        kursatkich = eng_kup_hudud(
            self.ustun, self.dan, self.gacha, self.hududlar, self.kitob)
        kursatkich2 = eng_katta(kursatkich, 1, urin=self.urin)
        kursatkich2 = self.hududlar_ru[kursatkich2]
        replace_text_in_doc(self.file_nomi, self.file_nomi,
                            self.uzgarmas, kursatkich2)

    @property
    def kursatkich(self):
        kursatkich = eng_kup_hudud(
            self.ustun, self.dan, self.gacha, self.hududlar, self.kitob)
        kursatkich2 = eng_katta(kursatkich, 0, urin=self.urin)
        replace_text_in_doc(self.file_nomi, self.file_nomi,
                            self.uzgarmas, kursatkich2)


def ustun_nomi(name, period):
    # Ko‘rsatkichlar ro‘yxati
    korsatkichlar = tarmoqlar.tarmoq

    # Yillar
    yillar = ['2023', '2024', '2025']

    # Lug'atni yaratamiz
    manzil = {}

    # A dan boshlab barcha ustunlarni belgilang
    start_index = 2  # A = 1

    for i, nom in enumerate(korsatkichlar):
        manzil[nom] = {}
        for j, yil in enumerate(yillar):
            col_index = start_index + i * len(yillar) + j
            col_letter = get_column_letter(col_index)
            manzil[nom][yil] = col_letter

    # Natijani chiqarish
    # pprint(manzil)
    return manzil[name][period]
