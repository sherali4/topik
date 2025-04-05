from func import replace_text_in_doc, eng_katta, eng_kup_hudud, from_excel, eng_kup, kiril_to_latin, Engkup
import openpyxl
from docx import Document
import re
from transliterate import translit


def from_excel(yacheyka):
    # Load the Excel file
    workbook = openpyxl.load_workbook('baza.xlsx', data_only=True)
    sheet = workbook.active

    # Get the value of the specified cell
    value = sheet[yacheyka].value
    workbook.close()
    value = str(value).replace('.', ',')
    return value

    # if value is not None:
    #    # Check if the value is a float, convert to string, and replace '.' with ','
    #    if isinstance(value, float):
    #        value = str(value).replace('.', ',')
    #    # Ensure that we don't try to replace in a non-string value (e.g., None or int)
    #    value_comma = str(value) if isinstance(value, str) else str(value)
    #    print(f"The value in cell {yacheyka} is: {value_comma}")
    #    return value_comma
    # else:
    #    print(f"No value found in cell {yacheyka}")
    #    return ''  # Return an empty string if no value is found


# def from_excel(yacheyka):
#
#    # Load the Excel file
#    workbook = openpyxl.load_workbook('baza.xlsx', data_only=True)
#
#    # Select the active sheet or a specific sheet
#    # or workbook['SheetName'] if you know the sheet name
#    sheet = workbook.active
#
#    # Get the value of cell B4
#    value_b4 = sheet[yacheyka].value
#    workbook.close()
#    value_b4_comma = str(value_b4).replace('.', ',')
#
#    # Print the value
#    print(f"The value in cell B4 is: {value_b4_comma}")
#    return str(value_b4).replace('.', ',') if value_b4 else ''


def replace_text_in_doc(doc_path, file_path, old_text, new_text):
    try:
        doc = Document(doc_path)

        for para in doc.paragraphs:
            for run in para.runs:
                run.text = run.text.replace(old_text, new_text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.text = run.text.replace(old_text, new_text)

        for section in doc.sections:
            for para in section.header.paragraphs + section.footer.paragraphs:
                for run in para.runs:
                    run.text = run.text.replace(old_text, new_text)

        doc.save(file_path)
        print(
            f"'{old_text}' -> '{new_text}'")
    except Exception as e:
        print(f"Xatolik yuz berdi: {e}")


def eng_kup_hudud(column, start_row, end_row, new_keys):
    workbook = openpyxl.load_workbook('baza.xlsx', data_only=True)
    sheet = workbook.active

    data_dict = {}

    # Excel qatorlarini yangi kalitlar bilan moslashtirish
    for i, row in enumerate(range(start_row, end_row + 1)):
        if i >= len(new_keys):  # Agar kalitlar ro‘yxati tugagan bo‘lsa, boshqa kalitni ishlatmasin
            break
        cell_value = sheet[f"{column}{row}"].value

        try:
            numeric_value = str(cell_value).replace('.', ',')
            data_dict[new_keys[i]] = numeric_value
        except ValueError:
            data_dict[new_keys[i]] = 0

    workbook.close()

    # Qiymatlar bo‘yicha Z-A (kamayish tartibida) saralash
    sorted_dict = dict(
        sorted(data_dict.items(), key=lambda item: item[1], reverse=True))

    return sorted_dict


new_keys = ['Andijon viloyati', 'Buxoro viloyati', 'Jizzax viloyati', 'Qashqadaryo viloyati', 'Navoiy viloyati', 'Namangan viloyati',
            'Samarqand viloyati', 'Surxondaryo viloyati', 'Sirdaryo viloyati', 'Toshkent shahri', 'Toshkent viloyati', "Farg'ona viloyati", 'Xorazm viloyati', "Qoraqalpog'iston Respublikasi"]


def eng_katta(data, val, urin):
    urin = urin - 1
    first_key = list(data.keys())[urin]
    first_value = data[first_key]
    if val == 1:
        return first_key
    else:
        return first_value


def eng_kup(ustun, dan, gacha, new_keys, urin, hudud=0):
    # eng_katta(eng_kup_hudud("B", 5, 18, new_keys), 1, 1)
    eng_katta(eng_kup_hudud(ustun, dan, gacha, new_keys), hudud, urin)


def kiril_to_latin(name):
    if isinstance(name[0], str) and re.search(r'[а-яА-Я]', name[0]):
        latin_name = translit(name, 'ru', reversed=True).upper()
        return latin_name
    else:
        return name.upper()


class Engkup:

    def __init__(self, file_nomi, uzgarmas, urin):
        # Private variablega o‘zgartirish kiritamiz
        self.file_nomi = file_nomi
        self.uzgarmas = uzgarmas
        self.urin = urin
        self.dan = 5
        self.gacha = self.dan + 13
        self.hududlar = ['Andijon viloyati', 'Buxoro viloyati', 'Jizzax viloyati', 'Qashqadaryo viloyati', 'Navoiy viloyati', 'Namangan viloyati',
                         'Samarqand viloyati', 'Surxondaryo viloyati', 'Sirdaryo viloyati', 'Toshkent shahri', 'Toshkent viloyati', "Farg'ona viloyati", 'Xorazm viloyati', "Qoraqalpog'iston Respublikasi"]
        self.hududlar_ru = {
            'Andijon viloyati': 'Андижанская область',
            'Buxoro viloyati': 'Бухарская область',
            'Jizzax viloyati': 'Джизакская область',
            'Qashqadaryo viloyati': 'Кашкадарьинская область',
            'Navoiy viloyati': 'Навоийская область',
            'Namangan viloyati': 'Наманганская область',
            'Samarqand viloyati': 'Самаркандская область',
            'Surxondaryo viloyati': 'Сурхандарьинская область',
            'Sirdaryo viloyati': 'Сирдарьинская область',
            'Toshkent shahri': 'Город Ташкент',
            'Toshkent viloyati': 'Ташкент область',
            "Farg'ona viloyati": "Ферганская область",
            'Xorazm viloyati': 'Хорезмская область',
            "Qoraqalpog'iston Respublikasi": "Республика Каракалпакстан"
        }

    @property
    def ustun(self):
        return self.ustun
    
    @property
    def 

    @ustun.setter
    def ustun(self, value):
        self.ustun = kiril_to_latin(value)  # Transliterate to Latin

    @property
    def hudud(self):
        kursatkich = eng_kup_hudud(
            self.ustun, self.dan, self.gacha, self.hududlar)
        kursatkich2 = eng_katta(kursatkich, 1, urin=self.urin)
        replace_text_in_doc(self.file_nomi, self.file_nomi,
                            self.uzgarmas, kursatkich2)

    @property
    def hudud_ru(self):
        kursatkich = eng_kup_hudud(
            self.ustun, self.dan, self.gacha, self.hududlar)
        kursatkich2 = eng_katta(kursatkich, 1, urin=self.urin)
        kursatkich2 = self.hududlar_ru[kursatkich2]
        replace_text_in_doc(self.file_nomi, self.file_nomi,
                            self.uzgarmas, kursatkich2)

    @property
    def kursatkich(self):
        kursatkich = eng_kup_hudud(
            self.ustun, self.dan, self.gacha, self.hududlar)
        kursatkich2 = eng_katta(kursatkich, 0, urin=self.urin)
        replace_text_in_doc(self.file_nomi, self.file_nomi,
                            self.uzgarmas, kursatkich2)


# period
replace_text_in_doc("example.docx", "357.docx", "period_uz", "mart")
replace_text_in_doc("357.docx", "357.docx", "period_ru", "март")

# Foydalanish:
replace_text_in_doc("357.docx", "357.docx", "2023_raqam", from_excel('B4'))
replace_text_in_doc("357.docx", "357.docx", "2024_raqam", from_excel('C4'))


Engkup.ustun = 'C'
print(Engkup.ustun)
Engkup('357.docx', '@h1', 1).hudud
Engkup('357.docx', '@hr1', 1).hudud_ru
Engkup('357.docx', '@k1', 1).kursatkich
Engkup('357.docx', '@h2', 2).hudud
Engkup('357.docx', '@hr2', 2).hudud_ru
Engkup('357.docx', '@k2', 2).kursatkich
