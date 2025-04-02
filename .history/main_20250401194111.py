from docx import Document
import openpyxl


def from_excel(yacheyka):

    # Load the Excel file
    workbook = openpyxl.load_workbook('baza.xlsx')

    # Select the active sheet or a specific sheet
    # or workbook['SheetName'] if you know the sheet name
    sheet = workbook.active

    # Get the value of cell B4
    value_b4 = sheet[yacheyka].value
    value_b4_comma = str(value_b4).replace('.', ',')

    # Print the value
    print(f"The value in cell B4 is: {value_b4_comma}")
    return value_b4_comma


def replace_text_in_doc(doc_path, file_path, old_text, new_text):
    doc = Document(doc_path)

    # Paragraflar ichidagi matnni almashtirish (formatni saqlab qolish)
    for para in doc.paragraphs:
        if old_text in para.text:
            # Har bir runni tekshiramiz va almashtiramiz
            for run in para.runs:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)

    # Jadval ichidagi matnni almashtirish (formatni saqlab qolish)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if old_text in cell.text:
                    # Har bir runni tekshiramiz va almashtiramiz
                    # Har bir paragraf uchun
                    for run in cell.paragraphs[0].runs:
                        if old_text in run.text:
                            run.text = run.text.replace(old_text, new_text)

    # Sarlavha va pastki qismni tekshirish (formatni saqlab qolish)
    sections = doc.sections
    for section in sections:
        # Header (sarlavha)
        for para in section.header.paragraphs:
            for run in para.runs:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)

        # Footer (pastki qism)
        for para in section.footer.paragraphs:
            for run in para.runs:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)

    # O‘zgarishlarni saqlash
    doc.save(file_path)
    print(f"'{old_text}' so‘zi '{new_text}' ga almashtirildi va '{file_path}' fayl saqlandi!")


# Foydalanish:
replace_text_in_doc("example.docx", "357.docx", "2025_raqam", "100")
replace_text_in_doc("357.docx", "357.docx", "2023_raqam", from_excel('B4'))
replace_text_in_doc("357.docx", "357.docx", "2024_raqam", from_excel('C4'))
