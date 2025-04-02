import win32com.client
from docx import Document
import openpyxl


def from_excel(yacheyka):

    # Load the Excel file
    workbook = openpyxl.load_workbook('baza.xlsx', data_only=True)

    # Select the active sheet or a specific sheet
    # or workbook['SheetName'] if you know the sheet name
    sheet = workbook.active

    # Get the value of cell B4
    value_b4 = sheet[yacheyka].value
    workbook.close()
    value_b4_comma = str(value_b4).replace('.', ',')

    # Print the value
    print(f"The value in cell B4 is: {value_b4_comma}")
    return str(value_b4).replace('.', ',') if value_b4 else ''


def replace_text_in_doc(doc_path, file_path, old_text, new_text):
    try:
        doc = Document(doc_path)

        for para in doc.paragraphs:
            for run in para.runs:
                run.text = run.text.replace(old_text, new_text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.text = run.text.replace(old_text, new_text)

        for section in doc.sections:
            for para in section.header.paragraphs + section.footer.paragraphs:
                for run in para.runs:
                    run.text = run.text.replace(old_text, new_text)

        doc.save(file_path)
        print(
            f"'{old_text}' -> '{new_text}' almashtirildi va '{file_path}' saqlandi!")
    except Exception as e:
        print(f"Xatolik yuz berdi: {e}")


# period
replace_text_in_doc("example.docx", "357.docx", "period_uz", "mart")
replace_text_in_doc("357.docx", "357.docx", "period_ru", "март")


# Foydalanish:
replace_text_in_doc("357.docx", "357.docx", "2025_raqam", "100")
replace_text_in_doc("357.docx", "357.docx", "2023_raqam", from_excel('B4'))
replace_text_in_doc("357.docx", "357.docx", "2024_raqam", from_excel('C4'))


def convert_docx_to_pdf(docx_path, pdf_path):
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False  # Word oynasi ko'rinmasligi uchun
    word.DisplayAlerts = False  # Ogohlantirishlarni o‘chirish

    try:
        doc = word.Documents.Open(docx_path)
        doc.SaveAs(pdf_path, FileFormat=17)  # PDF formatiga saqlash
        doc.Close(True)
    except Exception as e:
        print(f"Xatolik yuz berdi: {e}")
    finally:
        word.Quit()


# Example usage
convert_docx_to_pdf("357.docx", "357.pdf")
print("Conversion completed!")
